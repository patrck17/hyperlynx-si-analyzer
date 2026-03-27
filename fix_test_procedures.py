#!/usr/bin/env python3
"""
Fix specification values in tempvt.docx test procedure sections and rebuild
spec tables to match the format used in the Appendix B results tables.

Sections affected:
  1.3.5  EIA-422:    rise/fall → ≤ 50ns, rebuild table to 5-col RS-422 format
  1.3.7  EIA-485:    remove idle bus & noise margin sections entirely,
                     rebuild table to 5-col RS-485 format
  1.3.8  1.8V LVCMOS: fix VOH/VOL/NM/tr/tf specs, rebuild table to 4-col LVCMOS format
  1.3.11 3.3V LVCMOS: fix VOH/VOL/NM/tr/tf specs, rebuild table to 4-col LVCMOS format
"""

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Text replacement helpers
# ---------------------------------------------------------------------------

def replace_in_runs(paragraph, old_text, new_text):
    full = paragraph.text
    if old_text not in full:
        return False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    new_full = full.replace(old_text, new_text)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


# ---------------------------------------------------------------------------
# Paragraph creation helpers
# ---------------------------------------------------------------------------

def make_styled_paragraph(text, style_id):
    """Create a paragraph with a given style (e.g. Heading5, ListParagraph)."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def insert_noise_margin_test_step(doc, section_heading_text, table_ref, spec_text):
    """Insert a 'Noise Margin Verification' Heading 5 + list paragraphs
    before the table caption in the given test procedure section.

    Finds the Caption paragraph for the spec table and inserts just before it.
    """
    body = doc.element.body
    ns = WML_NS

    # Find the section heading
    section_start = None
    for i, elem in enumerate(body):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        ).strip()
        if section_heading_text in full_text:
            pStyle = elem.find(f'.//{{{ns}}}pStyle')
            if pStyle is not None and 'Heading3' in pStyle.get(f'{{{ns}}}val', ''):
                section_start = i
                break

    if section_start is None:
        return 0

    # Find the Caption paragraph (table reference) after the section heading
    caption_idx = None
    for i in range(section_start + 1, min(section_start + 60, len(body))):
        elem = body[i]
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        if pStyle is not None and 'Caption' in pStyle.get(f'{{{ns}}}val', ''):
            full_text = ''.join(
                t.text for t in elem.iter(f'{{{ns}}}t') if t.text
            ).strip()
            if table_ref in full_text:
                caption_idx = i
                break

    if caption_idx is None:
        return 0

    # Check if noise margin step already exists
    for i in range(section_start + 1, caption_idx):
        elem = body[i]
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        )
        if 'Noise Margin Verification' in full_text:
            return 0  # already present

    # Insert before the caption: Heading5 + 3 list paragraphs
    elems = [
        make_styled_paragraph("Noise Margin Verification:", "Heading5"),
        make_styled_paragraph(
            "Calculate the noise margin as the minimum of |VOD_high| and "
            "|VOD_low| minus the receiver input sensitivity threshold (200mV "
            f"per {spec_text}).",
            "ListParagraph"),
        make_styled_paragraph(
            f"Record the noise margin in {table_ref}.",
            "ListParagraph"),
        make_styled_paragraph(
            "Verify that the noise margin is \u2265 0.2V.",
            "ListParagraph"),
    ]

    for j, elem in enumerate(elems):
        body.insert(caption_idx + j, elem)

    return len(elems)


# ---------------------------------------------------------------------------
# Table creation (same pattern as insert_*_into_tempvt.py)
# ---------------------------------------------------------------------------

def make_table(doc, headers, data_rows, table_style="Table Connector Pinout"):
    """Create a detached table element matching the document's table style."""
    nrows = 1 + len(data_rows)
    ncols = len(headers)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(h)
        run.bold = True

    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            para = cell.paragraphs[0]
            para.clear()
            para.add_run(str(val))

    body = doc.element.body
    tbl_elem = table._tbl
    body.remove(tbl_elem)
    last_elem = body[-1]
    if (last_elem.tag.endswith('}p') and
            len(list(last_elem.iter(f'{{{WML_NS}}}t'))) == 0):
        body.remove(last_elem)
    return tbl_elem


# ---------------------------------------------------------------------------
# Find and replace a table in the body
# ---------------------------------------------------------------------------

def find_table_by_content(doc, marker_texts):
    """Find a table that contains ALL marker_texts in its cells.
    marker_texts should be a list of strings that must all be present.
    Returns (table_element, body_index)."""
    body = doc.element.body
    for i, elem in enumerate(body):
        if not elem.tag.endswith('}tbl'):
            continue
        full = ''.join(t.text for t in elem.iter(f'{{{WML_NS}}}t') if t.text)
        if all(m in full for m in marker_texts):
            return elem, i
    return None, None


def replace_table_at(doc, old_tbl_elem, body_idx, new_tbl_elem):
    """Replace an existing table element in the body at the given index."""
    body = doc.element.body
    body.remove(old_tbl_elem)
    # After removal, insert at the same position
    if body_idx >= len(body):
        body.append(new_tbl_elem)
    else:
        body.insert(body_idx, new_tbl_elem)


# ---------------------------------------------------------------------------
# Spec table definitions (matching Appendix B results formats)
# ---------------------------------------------------------------------------

RS422_SPEC_ROWS = [
    # (Parameter, P Leg, N Leg, Specification, Pass/Fail)
    ("Rise Time (ns)", "", "", "\u2264 50ns", ""),
    ("Fall Time (ns)", "", "", "\u2264 50ns", ""),
    ("Peak Voltage (V)", "", "", "\u2264 8V", ""),
    ("Trough Voltage (V)", "", "", "\u2265 -4V", ""),
    ("Overshoot (%)", "", "", "< 10% of amplitude", ""),
    ("Undershoot (%)", "", "", "< 10% of amplitude", ""),
    ("VOD (VP \u2013 VN) (V)", "", "", "\u2265 2.0V", ""),
    ("Noise Margin (V)", "", "", "\u2265 0.2V", ""),
]

RS485_SPEC_ROWS = [
    ("Rise Time (ns)", "", "", "\u2264 30ns", ""),
    ("Fall Time (ns)", "", "", "\u2264 30ns", ""),
    ("Peak Voltage (V)", "", "", "\u2264 8V", ""),
    ("Trough Voltage (V)", "", "", "\u2265 -4V", ""),
    ("Overshoot (%)", "", "", "< 10% of amplitude", ""),
    ("Undershoot (%)", "", "", "< 10% of amplitude", ""),
    ("VOD (VP \u2013 VN) (V)", "", "", "\u2265 1.5V", ""),
    ("Noise Margin (V)", "", "", "\u2265 0.2V", ""),
]

LVDS_SPEC_ROWS = [
    # (Parameter, P Leg, N Leg, Specification, Pass/Fail)
    ("Rise Time (ns)", "", "", "\u2264 2.08ns", ""),
    ("Fall Time (ns)", "", "", "\u2264 2.08ns", ""),
    ("Peak Voltage (V)", "", "", "\u2264 2.4V", ""),
    ("Trough Voltage (V)", "", "", "\u2265 0.0V", ""),
    ("Overshoot (%)", "", "", "< 10%", ""),
    ("Undershoot (%)", "", "", "< 10%", ""),
    ("VOD (VP \u2013 VN) (mV)", "", "", "250\u2013450mV", ""),
    ("VOS (V)", "", "", "1.125\u20131.375V", ""),
    ("Noise Margin (V)", "", "", "\u2265 0.1V", ""),
]

LVCMOS18_SPEC_ROWS = [
    # (Parameter, Measured Value, Specification, Pass/Fail)
    ("VOH (V)", "", "\u2265 1.35V", ""),
    ("VOL (V)", "", "\u2264 0.45V", ""),
    ("Rise Time (ns)", "", "\u2264 10ns", ""),
    ("Fall Time (ns)", "", "\u2264 10ns", ""),
    ("NMH (V)", "", "\u2265 0.18V", ""),
    ("NML (V)", "", "\u2265 0.18V", ""),
    ("Overshoot (V)", "", "Peak \u2264 2.1V", ""),
    ("Undershoot (V)", "", "Trough \u2265 \u22120.3V", ""),
]

LVCMOS33_SPEC_ROWS = [
    ("VOH (V)", "", "\u2265 2.4V", ""),
    ("VOL (V)", "", "\u2264 0.4V", ""),
    ("Rise Time (ns)", "", "\u2264 10ns", ""),
    ("Fall Time (ns)", "", "\u2264 10ns", ""),
    ("NMH (V)", "", "\u2265 0.4V", ""),
    ("NML (V)", "", "\u2265 0.4V", ""),
    ("Overshoot (V)", "", "Peak \u2264 3.6V", ""),
    ("Undershoot (V)", "", "Trough \u2265 \u22120.3V", ""),
]


# ---------------------------------------------------------------------------
# EIA-422 fixes
# ---------------------------------------------------------------------------

def fix_eia422_section(doc):
    """Fix EIA-422 (section 1.3.5): paragraph specs + replace spec table."""
    changes = 0

    # Fix paragraph text: rise/fall "10ns to 30ns" → "50ns"
    for para in doc.paragraphs:
        if "10ns to 30ns" in para.text:
            replace_in_runs(para, "10ns to 30ns", "50ns")
            changes += 1

    # Replace Table 18 (RS-422 spec table) — find by unique content
    old_tbl, idx = find_table_by_content(doc, ["Logic High", "RS-422"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["10–30", "Overshoot"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["RS-422 Specification"])

    if old_tbl is not None:
        headers = ["Parameter", "P Leg", "N Leg", "Specification", "Pass/Fail"]
        new_tbl = make_table(doc, headers, RS422_SPEC_ROWS)
        replace_table_at(doc, old_tbl, idx, new_tbl)
        changes += 1
        print(f"  EIA-422 (1.3.5): replaced spec table with 5-col format")
    else:
        print(f"  EIA-422 (1.3.5): spec table not found")

    # Add Noise Margin Verification test step
    added = insert_noise_margin_test_step(
        doc, "EIA-422 Signal Test", "Table 15", "TIA/EIA-422-B")
    changes += added

    print(f"  EIA-422 (1.3.5): {changes} corrections applied")
    return changes


# ---------------------------------------------------------------------------
# EIA-485 fixes
# ---------------------------------------------------------------------------

def remove_paragraphs_by_heading5(doc, heading_texts_to_remove,
                                   within_section=None):
    """Remove Heading 5 sections (heading + subsequent list paragraphs / normal
    paragraphs until the next heading or table caption).

    If within_section is given, only remove within that Heading 3 section.
    Deletes the actual XML elements from the body so no empty paragraphs remain.
    """
    body = doc.element.body
    ns = WML_NS
    removed = 0

    # Determine search scope
    start_idx = 0
    end_idx = len(body)
    if within_section:
        in_section = False
        for i, elem in enumerate(body):
            tag = elem.tag.split('}')[-1]
            if tag != 'p':
                continue
            pStyle = elem.find(f'.//{{{ns}}}pStyle')
            if pStyle is None:
                continue
            style_val = pStyle.get(f'{{{ns}}}val', '')
            full_text = ''.join(
                t.text for t in elem.iter(f'{{{ns}}}t') if t.text
            ).strip()
            if 'Heading3' in style_val and within_section in full_text:
                start_idx = i
                in_section = True
                continue
            if in_section and 'Heading3' in style_val:
                end_idx = i
                break

    # Collect paragraph elements that are in sections to remove
    elems_to_remove = []
    in_target_section = False

    for i, elem in enumerate(body):
        if i < start_idx or i >= end_idx:
            continue
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            if in_target_section:
                in_target_section = False
            continue

        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        style_val = pStyle.get(f'{{{ns}}}val') if pStyle is not None else None

        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        ).strip()

        if style_val and 'Heading5' in style_val:
            if any(ht in full_text for ht in heading_texts_to_remove):
                in_target_section = True
                elems_to_remove.append(elem)
                continue
            else:
                in_target_section = False
                continue

        if style_val and 'Heading' in style_val:
            in_target_section = False
            continue

        if style_val and 'Caption' in style_val:
            in_target_section = False
            continue

        if in_target_section:
            elems_to_remove.append(elem)

    for elem in elems_to_remove:
        body.remove(elem)
        removed += 1

    return removed


def fix_eia485_section(doc):
    """Fix EIA-485 (section 1.3.7): remove idle bus & noise margin sections,
    replace spec table with 5-col format."""
    changes = 0

    # Remove Idle Bus Verification and old Noise Margin Verification sections
    # (scoped to EIA-485 section only, so we don't remove EIA-422's noise margin)
    removed = remove_paragraphs_by_heading5(
        doc, ["Idle Bus Verification", "Noise Margin Verification"],
        within_section="EIA-485 Signal Test")
    changes += removed

    # Replace Table 20 (RS-485 spec table) — find by unique content
    old_tbl, idx = find_table_by_content(doc, ["+1.5V", "Differential Voltage"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["Idle Bus", "Rise Time"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["Differential Voltage (VH)"])

    if old_tbl is not None:
        headers = ["Parameter", "P Leg", "N Leg", "Specification", "Pass/Fail"]
        new_tbl = make_table(doc, headers, RS485_SPEC_ROWS)
        replace_table_at(doc, old_tbl, idx, new_tbl)
        changes += 1
        print(f"  EIA-485 (1.3.7): replaced spec table with 5-col format")
    else:
        print(f"  EIA-485 (1.3.7): spec table not found")

    # Add new Noise Margin Verification test step (correct definition)
    added = insert_noise_margin_test_step(
        doc, "EIA-485 Signal Test", "Table 17", "TIA/EIA-485")
    changes += added

    print(f"  EIA-485 (1.3.7): {changes} corrections (incl {removed} paragraphs removed)")
    return changes


# ---------------------------------------------------------------------------
# 1.8V LVCMOS fixes
# ---------------------------------------------------------------------------

def fix_lvcmos18_section(doc):
    """Fix 1.8V LVCMOS (section 1.3.8) spec values + replace spec table."""
    changes = 0

    for para in doc.paragraphs:
        text = para.text

        if "1.17V" in text and "VOH" in text:
            replace_in_runs(para, "1.17V", "1.35V")
            changes += 1
        if "0.63V" in text and "VOL" in text:
            replace_in_runs(para, "0.63V", "0.45V")
            changes += 1
        if "65% of VDD" in text:
            replace_in_runs(para, "65% of VDD", "75% of VDD")
            changes += 1
        if "35% of VDD" in text:
            replace_in_runs(para, "35% of VDD", "25% of VDD")
            changes += 1
        if "≤ 4ns" in text and ("rise" in text.lower() or "fall" in text.lower()):
            replace_in_runs(para, "4ns", "10ns")
            changes += 1
        if "0.27V" in text and "noise margin" in text.lower():
            replace_in_runs(para, "0.27V", "0.18V")
            changes += 1
        if "JESD76" in text:
            replace_in_runs(para, "JESD76", "JESD8-7A")
            changes += 1

    # Replace Table 21 (1.8V LVCMOS spec table) — find by unique content
    old_tbl, idx = find_table_by_content(doc, ["1.17", "VOH"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["0.63", "VOL"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["0.27", "Noise Margin"])

    if old_tbl is not None:
        headers = ["Parameter", "Measured Value", "Specification", "Pass/Fail"]
        new_tbl = make_table(doc, headers, LVCMOS18_SPEC_ROWS)
        replace_table_at(doc, old_tbl, idx, new_tbl)
        changes += 1
        print(f"  1.8V LVCMOS (1.3.8): replaced spec table with corrected format")
    else:
        print(f"  1.8V LVCMOS (1.3.8): spec table not found")

    print(f"  1.8V LVCMOS (1.3.8): {changes} corrections applied")
    return changes


# ---------------------------------------------------------------------------
# 3.3V LVCMOS fixes
# ---------------------------------------------------------------------------

def fix_lvcmos33_section(doc):
    """Fix 3.3V LVCMOS (section 1.3.11) spec values + replace spec table."""
    changes = 0

    for para in doc.paragraphs:
        text = para.text

        if "2.31V" in text and "VOH" in text:
            replace_in_runs(para, "2.31V", "2.4V")
            changes += 1
        if "0.99V" in text and "VOL" in text:
            replace_in_runs(para, "0.99V", "0.4V")
            changes += 1
        if "70% of VDD" in text:
            replace_in_runs(para, "70% of VDD", "73% of VDD")
            changes += 1
        if "30% of VDD" in text:
            replace_in_runs(para, "30% of VDD", "12% of VDD")
            changes += 1
        if "≤ 4ns" in text and ("rise" in text.lower() or "fall" in text.lower()):
            replace_in_runs(para, "4ns", "10ns")
            changes += 1
        if "0.533V" in text and "noise margin" in text.lower():
            replace_in_runs(para, "0.533V", "0.4V")
            changes += 1
        if "0.533 V" in text:
            replace_in_runs(para, "0.533 V", "0.4 V")
            changes += 1

    # Replace Table 24 (3.3V LVCMOS spec table) — find by unique content
    old_tbl, idx = find_table_by_content(doc, ["2.31", "VOH"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["0.99", "VOL"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["0.533", "Noise Margin"])

    if old_tbl is not None:
        headers = ["Parameter", "Measured Value", "Specification", "Pass/Fail"]
        new_tbl = make_table(doc, headers, LVCMOS33_SPEC_ROWS)
        replace_table_at(doc, old_tbl, idx, new_tbl)
        changes += 1
        print(f"  3.3V LVCMOS (1.3.11): replaced spec table with corrected format")
    else:
        print(f"  3.3V LVCMOS (1.3.11): spec table not found")

    print(f"  3.3V LVCMOS (1.3.11): {changes} corrections applied")
    return changes


# ---------------------------------------------------------------------------
# LVDS fixes
# ---------------------------------------------------------------------------

def fix_lvds_section(doc):
    """Fix LVDS (section 1.3.10) spec values + replace spec table."""
    changes = 0

    # Fix paragraph text: update specs to EIA-644 at 125MHz
    for para in doc.paragraphs:
        text = para.text
        # Fix rise/fall time specs
        if "260ps" in text and ("rise" in text.lower() or "fall" in text.lower()
                                or "transition" in text.lower()):
            replace_in_runs(para, "260ps", "2.08ns")
            changes += 1
        # Fix VOD range
        if "247" in text and "454" in text:
            replace_in_runs(para, "247", "250")
            replace_in_runs(para, "454", "450")
            changes += 1

    # Replace Table 23 (LVDS spec table) — find by unique content
    old_tbl, idx = find_table_by_content(doc, ["VOD", "LVDS"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["VOD", "Differential"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["247", "454"])
    if old_tbl is None:
        old_tbl, idx = find_table_by_content(doc, ["260", "Rise"])

    if old_tbl is not None:
        headers = ["Parameter", "P Leg", "N Leg", "Specification", "Pass/Fail"]
        new_tbl = make_table(doc, headers, LVDS_SPEC_ROWS)
        replace_table_at(doc, old_tbl, idx, new_tbl)
        changes += 1
        print(f"  LVDS (1.3.10): replaced spec table with 5-col EIA-644 format")
    else:
        print(f"  LVDS (1.3.10): spec table not found")

    # Add Noise Margin Verification test step
    added = insert_noise_margin_test_step(
        doc, "LVDS Signal Test", "Table 20", "EIA-644")
    changes += added

    print(f"  LVDS (1.3.10): {changes} corrections applied")
    return changes


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def fix_all(doc):
    """Fix all test procedure spec values and tables."""
    print("Fixing test procedure specifications...")
    total = 0
    total += fix_eia422_section(doc)
    total += fix_eia485_section(doc)
    total += fix_lvds_section(doc)
    total += fix_lvcmos18_section(doc)
    total += fix_lvcmos33_section(doc)
    print(f"  Total: {total} corrections applied")
    return total


if __name__ == "__main__":
    print("Loading tempvt.docx...")
    doc = Document("tempvt.docx")
    fix_all(doc)
    doc.save("tempvt.docx")
    print("Saved tempvt.docx")
