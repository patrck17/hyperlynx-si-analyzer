#!/usr/bin/env python3
"""
Insert RS-422 SI analysis results into tempvt.docx Appendix B –
"EIA-422 Signal Test" section.

Each differential pair gets a Heading 3 subsection (e.g., CAL_SPI_SCLK_P/N)
with a 4-column parameter table matching the existing format:
  Parameter | Measured Value | RS-422 Specification | Pass/Fail

Usage:
  python insert_rs422_into_tempvt.py                    # all 15 pairs
  python insert_rs422_into_tempvt.py --test 2           # first 2 pairs only
  python insert_rs422_into_tempvt.py --test 2 -o test_rs422.docx
"""

import csv
import re
import copy
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_parens(spec_str):
    """Remove parenthetical text from a spec string, preserving footnote markers."""
    return re.sub(r'\s*\([^)]*\)', '', spec_str).strip()

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# TIA/EIA-422-B specifications
RS422_SPECS = {
    "vod_min": 2.0,         # |VOD| ≥ 2.0V at receiver (across Rt)
    "tr_tf_max": 50.0,      # transition time ≤ 50% of unit interval (100ns @ 10Mbps)
    "overshoot_max_pct": 10.0,  # < 10% of signal amplitude
    "se_peak_max": 8.0,     # single-ended peak ≤ 8V
    "se_trough_min": -4.0,  # single-ended trough ≥ -4V
    "vod_diff_max": 12.0,   # |VOD| ≤ 12V differential
    "vid_threshold": 0.2,   # receiver input sensitivity ±200mV (TIA/EIA-422-B)
}


def load_rs422_csv(csv_path):
    """Load RS-422 results CSV and group by differential pair.

    Returns OrderedDict: pair_name → {"P": row_dict, "N": row_dict}
    """
    pairs = OrderedDict()
    with open(csv_path, newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            pair = row["diff_pair_name"]
            leg = row["leg"]
            if pair not in pairs:
                pairs[pair] = {}
            pairs[pair][leg] = row
    return pairs


def pass_fail(value, spec, op):
    """Return 'PASS' or 'FAIL'."""
    if op == ">=":
        return "PASS" if value >= spec else "FAIL"
    elif op == "<=":
        return "PASS" if value <= spec else "FAIL"
    return "PASS"


def build_rs422_table_rows(p_data, n_data):
    """Build combined table rows for one differential pair.

    Table format: Parameter | P Leg | N Leg | Specification | Pass/Fail

    Single-ended parameters show independent P and N values.
    Differential parameters span the P/N columns (merged visually via "—").

    Returns list of (parameter, p_val, n_val, spec, pass_fail).
    """
    rows = []

    voh_p = float(p_data["voh"])
    vol_p = float(p_data["vol"])
    voh_n = float(n_data["voh"])
    vol_n = float(n_data["vol"])
    tr_p = float(p_data["rise_time_ns"])
    tf_p = float(p_data["fall_time_ns"])
    tr_n = float(n_data["rise_time_ns"])
    tf_n = float(n_data["fall_time_ns"])
    peak_p = float(p_data["peak_v"])
    peak_n = float(n_data["peak_v"])
    trough_p = float(p_data["trough_v"])
    trough_n = float(n_data["trough_v"])
    os_p = float(p_data["overshoot_pct"])
    os_n = float(n_data["overshoot_pct"])
    us_p = float(p_data["undershoot_pct"])
    us_n = float(n_data["undershoot_pct"])

    # --- Single-ended parameters (per-leg) ---
    rows.append((
        "Rise Time (ns)",
        f"{tr_p:.2f}", f"{tr_n:.2f}",
        "\u2264 50ns",
        pass_fail(max(tr_p, tr_n), RS422_SPECS["tr_tf_max"], "<="),
    ))
    rows.append((
        "Fall Time (ns)",
        f"{tf_p:.2f}", f"{tf_n:.2f}",
        "\u2264 50ns",
        pass_fail(max(tf_p, tf_n), RS422_SPECS["tr_tf_max"], "<="),
    ))
    rows.append((
        "Peak Voltage (V)",
        f"{peak_p:.2f}", f"{peak_n:.2f}",
        "\u2264 8V",
        pass_fail(max(peak_p, peak_n), RS422_SPECS["se_peak_max"], "<="),
    ))
    rows.append((
        "Trough Voltage (V)",
        f"{trough_p:.2f}", f"{trough_n:.2f}",
        "\u2265 -4V",
        pass_fail(min(trough_p, trough_n), RS422_SPECS["se_trough_min"], ">="),
    ))
    rows.append((
        "Overshoot (%)",
        f"{os_p:.2f}", f"{os_n:.2f}",
        "< 10% of amplitude",
        pass_fail(max(os_p, os_n), RS422_SPECS["overshoot_max_pct"], "<="),
    ))
    rows.append((
        "Undershoot (%)",
        f"{us_p:.2f}", f"{us_n:.2f}",
        "< 10% of amplitude",
        pass_fail(max(us_p, us_n), RS422_SPECS["overshoot_max_pct"], "<="),
    ))

    # --- Differential parameter ---
    # VOD = VP - VN: when P high/N low → positive, when P low/N high → negative
    # Full swing = VOD_high - VOD_low (e.g. 2.3 - (-2.3) = 4.6V)
    vod_h = voh_p - vol_n   # P high, N low
    vod_l = vol_p - voh_n   # P low, N high
    vod_swing = vod_h - vod_l  # full differential swing

    rows.append((
        "VOD (VP \u2013 VN) (V)",
        f"{vod_swing:.2f}", "",
        "\u2265 2.0V",
        pass_fail(vod_swing, RS422_SPECS["vod_min"], ">="),
    ))

    # Noise margin: min(|VOD_high|, |VOD_low|) - 200mV receiver threshold
    vid = RS422_SPECS["vid_threshold"]
    nm = min(abs(vod_h), abs(vod_l)) - vid
    rows.append((
        "Noise Margin (V)",
        f"{nm:.2f}", "",
        "\u2265 0.2V",
        pass_fail(nm, vid, ">="),
    ))

    return rows


def find_appendix_b_section(doc, section_heading_text):
    """Find the Heading 2 in Appendix B with the given text.
    Returns (body, heading_elem_index, next_heading2_index).
    """
    body = doc.element.body
    ns = WML_NS

    appendix_start = None
    for i, elem in enumerate(body):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        )
        if 'Appendix B' in full_text:
            pStyle = elem.find(f'.//{{{ns}}}pStyle')
            if pStyle is not None and 'Heading1' == pStyle.get(f'{{{ns}}}val'):
                appendix_start = i
                break

    if appendix_start is None:
        raise ValueError("Could not find 'Appendix B' Heading 1")

    heading_idx = None
    for i in range(appendix_start + 1, len(body)):
        elem = body[i]
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        if pStyle is None:
            continue
        style = pStyle.get(f'{{{ns}}}val')
        if style != 'Heading2':
            continue
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        )
        if full_text.strip() == section_heading_text.strip():
            heading_idx = i
            break

    if heading_idx is None:
        raise ValueError(
            f"Could not find Heading 2 '{section_heading_text}' in Appendix B"
        )

    next_heading_idx = None
    for i in range(heading_idx + 1, len(body)):
        elem = body[i]
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        if pStyle is None:
            continue
        style = pStyle.get(f'{{{ns}}}val')
        if style == 'Heading2':
            next_heading_idx = i
            break

    if next_heading_idx is None:
        next_heading_idx = len(body)

    return body, heading_idx, next_heading_idx


def clear_existing_content(body, heading_idx, next_heading_idx):
    """Remove existing content between the section heading and the next Heading 2."""
    to_remove = []
    for i in range(heading_idx + 1, next_heading_idx):
        to_remove.append(body[i])
    for elem in to_remove:
        body.remove(elem)
    return heading_idx + 1


def make_heading_paragraph(text, heading_style_id):
    """Create a new heading paragraph element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), heading_style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_empty_paragraph():
    return OxmlElement('w:p')


def make_table(doc, headers, data_rows, table_style="Table Connector Pinout"):
    """Create a 4-column table matching the existing RS-422 format."""
    nrows = 1 + len(data_rows)
    ncols = len(headers)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(h)
        run.bold = True

    # Data rows
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        for ci, cell_info in enumerate(row_data):
            cell = row.cells[ci]
            para = cell.paragraphs[0]
            para.clear()

            if isinstance(cell_info, dict):
                text = cell_info["text"]
                run = para.add_run(text)
                if cell_info.get("bold"):
                    run.bold = True
                if cell_info.get("color"):
                    run.font.color.rgb = cell_info["color"]
                if cell_info.get("align"):
                    para.alignment = cell_info["align"]
            else:
                run = para.add_run(str(cell_info))

    # Detach table from end of body
    body = doc.element.body
    tbl_elem = table._tbl
    body.remove(tbl_elem)
    last_elem = body[-1]
    if (last_elem.tag.endswith('}p') and
            len(list(last_elem.iter(f'{{{WML_NS}}}t'))) == 0):
        body.remove(last_elem)

    return tbl_elem


def make_italic_paragraph(text, size_half_pts=18):
    """Create an italic paragraph at the given size (half-points)."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_elem = OxmlElement('w:i')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_half_pts))
    rPr.append(i_elem)
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def format_pf_cell(pf):
    """Return a cell_info dict for a Pass/Fail value."""
    if pf == "FAIL":
        return {"text": pf, "bold": True,
                "color": RGBColor(0xCC, 0x00, 0x00),
                "align": WD_ALIGN_PARAGRAPH.CENTER}
    elif pf == "PASS":
        return {"text": pf, "bold": True,
                "color": RGBColor(0x00, 0x80, 0x00),
                "align": WD_ALIGN_PARAGRAPH.CENTER}
    return pf  # "—" or other


def insert_rs422_data(doc, csv_path, max_pairs=None):
    """Insert RS-422 differential pair tables into Appendix B.

    Table format: Parameter | P Leg | N Leg | Specification | Pass/Fail
    Single-ended rows show per-leg values, differential rows at bottom.
    """
    body, heading_idx, next_heading_idx = find_appendix_b_section(
        doc, "EIA-422 Signal Test")

    pairs = load_rs422_csv(csv_path)

    # Clear existing placeholder content
    insert_idx = clear_existing_content(body, heading_idx, next_heading_idx)

    headers = [
        "Parameter",
        {"text": "P Leg", "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
        {"text": "N Leg", "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
        "Specification",
        {"text": "Pass/Fail", "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
    ]
    header_texts = ["Parameter", "P Leg", "N Leg", "Specification", "Pass/Fail"]

    count = 0
    for pair_name, legs in pairs.items():
        if max_pairs is not None and count >= max_pairs:
            break

        p_data = legs.get("P")
        n_data = legs.get("N")
        if not p_data or not n_data:
            continue

        # Heading 3: pair name: driver pins → load
        p_pin = p_data['driver_pin']
        n_pin = n_data['driver_pin']
        heading_text = (f"{pair_name}: "
                        f"{p_pin} / {n_pin} \u2192 "
                        f"{p_data['load_label']} / {n_data['load_label']}")
        h_elem = make_heading_paragraph(heading_text, "Heading3")
        body.insert(insert_idx, h_elem)
        insert_idx += 1

        # Build parameter rows
        param_rows = build_rs422_table_rows(p_data, n_data)

        # Build table data (5 columns)
        table_data = []
        for param, p_val, n_val, spec, pf in param_rows:
            table_data.append([
                param,
                {"text": p_val, "align": WD_ALIGN_PARAGRAPH.CENTER},
                {"text": n_val, "align": WD_ALIGN_PARAGRAPH.CENTER},
                spec,
                format_pf_cell(pf),
            ])

        tbl_elem = make_table(doc, header_texts, table_data)
        body.insert(insert_idx, tbl_elem)
        insert_idx += 1

        # Spacer
        body.insert(insert_idx, make_empty_paragraph())
        insert_idx += 1

        count += 1

    # Footnote
    fn_text = ("RS-422 specifications per TIA/EIA-422-B. "
               "Transition time limit = 50% of unit interval (100ns at 10 Mbps). "
               "Load: 3ft twisted pair cable (Z0=120\u03A9, vf=66%) into "
               "124\u03A9 differential termination. "
               "Driver model: AM26LV31ESDREP "
               "(Rout=25\u03A9, tr/tf=10ns, datasheet estimate).")
    fn_elem = make_italic_paragraph(fn_text, 16)
    body.insert(insert_idx, fn_elem)
    insert_idx += 1

    print(f"  Inserted {count} differential pairs into "
          f"Appendix B 'EIA-422 Signal Test'")


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="Insert RS-422 SI results into tempvt.docx"
    )
    parser.add_argument("-o", "--output", default=None,
                        help="Output file (default: overwrite tempvt.docx)")
    parser.add_argument("--csv", default="rs422_results.csv",
                        help="RS-422 results CSV")
    parser.add_argument("--test", type=int, default=None,
                        help="Insert only first N pairs (for review)")

    args = parser.parse_args()

    print("Loading tempvt.docx...")
    doc = Document("tempvt.docx")

    print("\n=== Inserting RS-422 SI data ===")
    insert_rs422_data(doc, args.csv, max_pairs=args.test)

    output_path = args.output or "tempvt.docx"
    doc.save(output_path)
    print(f"\nSaved to {output_path}")


if __name__ == "__main__":
    main()
