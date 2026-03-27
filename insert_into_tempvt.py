#!/usr/bin/env python3
"""
Insert SI analysis results into tempvt.docx Appendix B – Results sections
for 1.8V and 3.3V LVCMOS Signal Test, maintaining the source document's
formatting.

Inserts per-signal subsections with driver→receiver sub-subsections (where
applicable) and 7-column parameter tables, using the document's existing
heading styles and table style.
"""

import csv
import re
import copy
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import re as _re

from update_docx import (
    JEDEC_SPECS, RS422_TR_TF_SPEC, RS422_FOOTNOTE,
    LVC_INPUT_TR_TF_SPEC, LVC_INPUT_FOOTNOTE,
    HVD75_MODELS,
    build_resistor_pin_map, resolve_pins_in_row, load_csv,
    is_hvd75_signal, get_tr_tf_override, pass_fail, build_param_rows,
)


def strip_parens(spec_str):
    """Remove parenthetical text from a spec string, preserving footnote markers.
    '≤ 10ns (JESD8C.01)' → '≤ 10ns'
    '≤ 50ns (TIA/EIA-422-B) †' → '≤ 50ns †'
    """
    return _re.sub(r'\s*\([^)]*\)', '', spec_str).strip()


def add_rx_spacing(rx_str):
    """Add space after commas in receiver pin strings where missing.
    'U128-4,U128-1' → 'U128-4, U128-1'
    """
    return _re.sub(r',(?! )', ', ', rx_str)


WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def find_appendix_b_section(doc, section_heading_text):
    """Find the Heading 2 in Appendix B with the given text.

    Returns (body, heading_elem_index, next_heading2_index).
    The next_heading2_index is where we insert before.
    """
    body = doc.element.body
    ns = WML_NS

    # First find "Appendix B" heading to limit search scope
    appendix_start = None
    for i, elem in enumerate(body):
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        )
        if 'Appendix B' in full_text:
            pStyle = elem.find(f'.//{{{ns}}}pStyle')
            if pStyle is not None and 'Heading1' == pStyle.get(f'{{{ns}}}val'):
                appendix_start = i
                break

    if appendix_start is None:
        raise ValueError("Could not find 'Appendix B' Heading 1")

    # Find the Heading 2 matching section_heading_text after Appendix B
    heading_idx = None
    for i in range(appendix_start + 1, len(body)):
        elem = body[i]
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        if pStyle is None:
            continue
        style = pStyle.get(f'{{{ns}}}val')
        if style != 'Heading2':
            continue
        full_text = ''.join(
            t.text for t in elem.iter(f'{{{ns}}}t') if t.text
        )
        if full_text.strip() == section_heading_text.strip():
            heading_idx = i
            break

    if heading_idx is None:
        raise ValueError(
            f"Could not find Heading 2 '{section_heading_text}' in Appendix B"
        )

    # Find the next Heading 2 (start of next results section)
    next_heading_idx = None
    for i in range(heading_idx + 1, len(body)):
        elem = body[i]
        tag = elem.tag.split('}')[-1]
        if tag != 'p':
            continue
        pStyle = elem.find(f'.//{{{ns}}}pStyle')
        if pStyle is None:
            continue
        style = pStyle.get(f'{{{ns}}}val')
        if style == 'Heading2':
            next_heading_idx = i
            break

    if next_heading_idx is None:
        # Last section — insert at end of body
        next_heading_idx = len(body)

    return body, heading_idx, next_heading_idx


def clear_existing_content(body, heading_idx, next_heading_idx):
    """Remove existing content between the section heading and the next
    Heading 2 (i.e., clear old placeholder data like SYNC_EN2_P3V3)."""
    # Remove elements from heading_idx+1 to next_heading_idx-1 (exclusive)
    to_remove = []
    for i in range(heading_idx + 1, next_heading_idx):
        to_remove.append(body[i])
    for elem in to_remove:
        body.remove(elem)
    # After removal, the next heading is now at heading_idx + 1
    return heading_idx + 1


def make_heading_paragraph(text, heading_style_id):
    """Create a new heading paragraph element with the given style ID."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), heading_style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def make_empty_paragraph():
    """Create an empty paragraph element."""
    return OxmlElement('w:p')


def make_table(doc, headers, data_rows, table_style="Table Connector Pinout"):
    """Create a table matching the source document's style.

    Returns the table's OxmlElement (detached from the document body).
    """
    nrows = 1 + len(data_rows)
    ncols = len(headers)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(h)
        run.bold = True

    # Data rows
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        for ci, cell_info in enumerate(row_data):
            cell = row.cells[ci]
            para = cell.paragraphs[0]
            para.clear()

            if isinstance(cell_info, dict):
                text = cell_info["text"]
                run = para.add_run(text)
                if cell_info.get("bold"):
                    run.bold = True
                if cell_info.get("color"):
                    run.font.color.rgb = cell_info["color"]
                if cell_info.get("align"):
                    para.alignment = cell_info["align"]
            else:
                run = para.add_run(str(cell_info))

    # Detach the table from the end of the body
    body = doc.element.body
    tbl_elem = table._tbl
    body.remove(tbl_elem)
    # Remove trailing empty paragraph left by doc.add_table()
    last_elem = body[-1]
    if (last_elem.tag.endswith('}p') and
            len(list(last_elem.iter(f'{{{WML_NS}}}t'))) == 0):
        body.remove(last_elem)

    return tbl_elem


def build_run_table_data(run_data, specs, vdd, r_net_map, signal_net_name):
    """Build 4-column table data for one run case.

    Returns (headers, data_rows, override, driver, rx).
    Driver/receiver info is returned separately for use in headings.
    """
    headers = ["Parameter", "Measured Value", "Specification", "Pass/Fail"]

    override = get_tr_tf_override(run_data)
    param_rows = build_param_rows(run_data, specs, vdd,
                                  tr_tf_override=override)

    driver_raw = run_data["driver_pin"]
    rx_raw = run_data["rx_pins"]
    driver, rx = resolve_pins_in_row(driver_raw, rx_raw, r_net_map,
                                     signal_net_name)
    rx = add_rx_spacing(rx)

    data_rows = []
    for param, meas, spec, pf in param_rows:
        pf_color = (RGBColor(0xCC, 0x00, 0x00) if pf == "FAIL"
                    else RGBColor(0x00, 0x80, 0x00))
        data_rows.append([
            param,
            {"text": meas, "align": WD_ALIGN_PARAGRAPH.CENTER},
            strip_parens(spec),
            {"text": pf, "bold": True, "color": pf_color,
             "align": WD_ALIGN_PARAGRAPH.CENTER},
        ])

    return headers, data_rows, override, driver, rx


def make_footnote_paragraph(text):
    """Create an italic 8pt footnote paragraph."""
    fn_p = OxmlElement('w:p')
    fn_r = OxmlElement('w:r')
    fn_rPr = OxmlElement('w:rPr')
    fn_i = OxmlElement('w:i')
    fn_sz = OxmlElement('w:sz')
    fn_sz.set(qn('w:val'), '16')  # 8pt = 16 half-points
    fn_rPr.append(fn_i)
    fn_rPr.append(fn_sz)
    fn_r.append(fn_rPr)
    fn_t = OxmlElement('w:t')
    fn_t.text = text
    fn_t.set(qn('xml:space'), 'preserve')
    fn_r.append(fn_t)
    fn_p.append(fn_r)
    return fn_p


# JEDEC section footnotes
JEDEC_FOOTNOTE = {
    3.3: "Specifications per JESD8C.01 (3.3V LVCMOS).",
    1.8: "Specifications per JESD8-7A (1.8V LVCMOS).",
}


def insert_si_data(doc, section_heading, csv_path, vdd, r_net_map,
                   max_signals=None):
    """Insert per-signal SI analysis tables into Appendix B after the
    given section heading.

    Appendix B uses:
      Heading 2 = section (e.g., "1.8V LVCMOS Signal Test")
      Heading 3 = signal name + driver → receiver
      Heading 4 = driver→receiver direction (for A/B runs only)
    Tables are 4 columns: Parameter | Measured Value | Specification | Pass/Fail

    If max_signals is set, only the first N signals are inserted (for testing).
    """
    body, heading_idx, next_heading_idx = find_appendix_b_section(
        doc, section_heading)
    specs = JEDEC_SPECS[vdd]
    signals = load_csv(csv_path)

    # Clear any existing placeholder content in this section
    insert_idx = clear_existing_content(body, heading_idx, next_heading_idx)

    has_rs422_footnote = False
    has_lvc_input_footnote = False
    sig_count = 0

    for sig_name, run_cases in signals.items():
        if max_signals is not None and sig_count >= max_signals:
            break
        sig_count += 1

        has_multiple_runs = (len(run_cases) > 1 and
                            any(rid in ("A", "B") for rid, _ in run_cases))

        if has_multiple_runs:
            # Signal heading (Heading 3) — just the signal name
            h_elem = make_heading_paragraph(sig_name, "Heading3")
            body.insert(insert_idx, h_elem)
            insert_idx += 1

            for run_id, run_data in run_cases:
                # Build table (also resolves pins)
                headers, data_rows, override, driver_res, rx_res = \
                    build_run_table_data(
                        run_data, specs, vdd, r_net_map, sig_name)

                # Direction heading (Heading 4): driver → receiver
                sub_text = f"{driver_res} \u2192 {rx_res}"
                h_sub = make_heading_paragraph(sub_text, "Heading4")
                body.insert(insert_idx, h_sub)
                insert_idx += 1

                # Table
                tbl_elem = make_table(doc, headers, data_rows)
                body.insert(insert_idx, tbl_elem)
                insert_idx += 1

                # Spacer
                body.insert(insert_idx, make_empty_paragraph())
                insert_idx += 1

                if override == "rs422":
                    has_rs422_footnote = True
                elif override == "lvc_input":
                    has_lvc_input_footnote = True
        else:
            _, run_data = run_cases[0]

            # Build table (also resolves pins)
            headers, data_rows, override, driver_res, rx_res = \
                build_run_table_data(
                    run_data, specs, vdd, r_net_map, sig_name)

            # Signal heading (Heading 3): signal name + driver → receiver
            heading_text = (f"{sig_name}: "
                            f"{driver_res} \u2192 {rx_res}")
            h_elem = make_heading_paragraph(heading_text, "Heading3")
            body.insert(insert_idx, h_elem)
            insert_idx += 1

            # Table
            tbl_elem = make_table(doc, headers, data_rows)
            body.insert(insert_idx, tbl_elem)
            insert_idx += 1

            # Spacer
            body.insert(insert_idx, make_empty_paragraph())
            insert_idx += 1

            if override == "rs422":
                has_rs422_footnote = True
            elif override == "lvc_input":
                has_lvc_input_footnote = True

    # Section footnotes
    body.insert(insert_idx, make_empty_paragraph())
    insert_idx += 1

    # JEDEC standard footnote
    body.insert(insert_idx,
                make_footnote_paragraph(JEDEC_FOOTNOTE[vdd]))
    insert_idx += 1

    # RS-422 / LVC input footnotes (when applicable)
    if has_rs422_footnote:
        body.insert(insert_idx, make_footnote_paragraph(RS422_FOOTNOTE))
        insert_idx += 1
    if has_lvc_input_footnote:
        body.insert(insert_idx, make_footnote_paragraph(LVC_INPUT_FOOTNOTE))
        insert_idx += 1

    print(f"  Inserted {sig_count} signals into "
          f"Appendix B '{section_heading}'")
    if has_rs422_footnote:
        print(f"    (RS-422 tr/tf: \u2264 50ns for LVC8T245 \u2192 SN55HVD75)")
    if has_lvc_input_footnote:
        print(f"    (LVC8T245 input tr/tf: \u2264 33ns for SN55HVD75 \u2192 LVC8T245)")


def main():
    print("Loading tempvt.docx...")
    doc = Document("tempvt.docx")
    r_net_map = build_resistor_pin_map("temp.qcv")

    # Insert P1V8 data into Appendix B "1.8V LVCMOS Signal Test"
    print("\n=== Inserting P1V8 SI data ===")
    insert_si_data(doc, "1.8V LVCMOS Signal Test", "p1v8_results.csv",
                   1.8, r_net_map)

    # Insert P3V3 data into Appendix B "3.3V LVCMOS Signal Test"
    print("\n=== Inserting P3V3 SI data ===")
    insert_si_data(doc, "3.3V LVCMOS Signal Test", "p3v3_results.csv",
                   3.3, r_net_map)

    output_path = "tempvt.docx"
    doc.save(output_path)
    print(f"\nSaved to {output_path}")


if __name__ == "__main__":
    main()
