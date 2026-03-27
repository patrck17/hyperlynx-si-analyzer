#!/usr/bin/env python3
"""
Rebuild Word document SI analysis reports with JEDEC-compliant specs.

Structure:
  3.7  3.3V LVCMOS Signal Test (JESD8C.01)
    3.7.1  SIGNAL_NAME_1
      3.7.1.1  Run A: U128-1 → U128-4      (when A/B runs exist)
        [7-row table]
      3.7.1.2  Run B: U126-3 → U128-4, U128-1
        [7-row table]
    3.7.2  SIGNAL_NAME_2                     (single-run signals)
      [7-row table]

Specs per JEDEC:
  JESD8C.01 (3.3V): VOH>=2.4V, VOL<=0.4V, VIH=2.0V, VIL=0.8V,
                     NM>=0.4V, tr/tf<=10ns, OS<=VDD+0.3V/>=-0.3V
  JESD8-7A  (1.8V): VOH>=1.35V, VOL<=0.45V, VIH=1.17V, VIL=0.63V,
                     NM>=0.18V, tr/tf<=10ns, OS<=VDD+0.3V/>=-0.3V

RS-422 exception (SN55HVD75 signals):
  tr/tf <= 50ns per TIA/EIA-422-B (50% of 100ns unit interval @ 10 Mbps)

Usage:
    python3 update_docx.py
"""

import csv
import re
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# JEDEC specifications
# ---------------------------------------------------------------------------

JEDEC_SPECS = {
    3.3: {
        "standard": "JESD8C.01",
        "voh": ("≥ 2.4V", 2.4, ">="),
        "vol": ("≤ 0.4V", 0.4, "<="),
        "tr":  ("≤ 10ns (JESD8C.01)", 10.0, "<="),
        "tf":  ("≤ 10ns (JESD8C.01)", 10.0, "<="),
        "nmh": ("≥ 0.4V", 0.4, ">="),
        "nml": ("≥ 0.4V", 0.4, ">="),
        "os_peak": 3.6,   # VDD + 0.3V
        "os_trough": -0.3,
        "os_spec": "Peak ≤ 3.6V, Trough ≥ −0.3V",
    },
    1.8: {
        "standard": "JESD8-7A",
        "voh": ("≥ 1.35V (VDD−0.45V)", 1.35, ">="),
        "vol": ("≤ 0.45V", 0.45, "<="),
        "tr":  ("≤ 10ns (JESD8-7A)", 10.0, "<="),
        "tf":  ("≤ 10ns (JESD8-7A)", 10.0, "<="),
        "nmh": ("≥ 0.18V", 0.18, ">="),
        "nml": ("≥ 0.18V", 0.18, ">="),
        "os_peak": 2.1,   # VDD + 0.3V
        "os_trough": -0.3,
        "os_spec": "Peak ≤ 2.1V, Trough ≥ −0.3V",
    },
}

# RS-422 rise/fall time spec for signals FROM LVC8T245 TO SN55HVD75
# TIA/EIA-422-B: transition time <= 50% of unit interval
# At 10 Mbps: unit interval = 100ns, max transition time = 50ns
RS422_TR_TF_SPEC = ("≤ 50ns (TIA/EIA-422-B) †", 50.0, "<=")

# LVC8T245 input rise/fall time spec for signals FROM SN55HVD75 TO LVC8T245
# SN74LVC8T245 datasheet: max input transition rate = 10ns/V
# At VCC = 3.3V: 10ns/V × 3.3V = 33ns
LVC_INPUT_TR_TF_SPEC = ("≤ 33ns (SN74LVC8T245, 10ns/V) ††", 33.0, "<=")

RS422_FOOTNOTE = (
    "† Rise/fall time (LVC8T245 → SN55HVD75 direction): JESD8C.01 specifies "
    "≤ 10ns for general-purpose 3.3V LVCMOS. These signals interface with "
    "SN55HVD75 RS-422 transceivers operating at ≤ 10 Mbps. Per TIA/EIA-422-B, "
    "the transition time shall not exceed 50% of the unit interval "
    "(100ns at 10 Mbps), yielding a relaxed limit of ≤ 50ns."
)

LVC_INPUT_FOOTNOTE = (
    "†† Rise/fall time (SN55HVD75 → LVC8T245 direction): The SN74LVC8T245 "
    "datasheet specifies a maximum input transition rate of 10ns/V. "
    "At VCC = 3.3V, the maximum allowable input rise/fall time is "
    "10 × 3.3 = 33ns."
)

# Models that indicate SN55HVD75 involvement
HVD75_MODELS = {"HVD75_R", "HVD75_D", "HVD75_EN"}
# HVD75 driver models (SN55 is the signal source)
HVD75_DRIVER_MODELS = {"HVD75_R"}


# ---------------------------------------------------------------------------
# Netlist pin resolution
# ---------------------------------------------------------------------------

def build_resistor_pin_map(qcv_path):
    """Parse temp.qcv to build a mapping: series_R_refdes -> far-side pin.

    For each resistor Rxxx, finds both nets it connects. Returns a dict:
      { "R426": { "TR0_EXT_SYNC_P3V3": ["U128-4", "U128-1"],
                  "UNNAMED_775": ["U126-3"] }, ... }
    """
    # Build: R_refdes -> list of (net_name, pin_side, other_pins_on_net)
    r_nets = {}  # R_refdes -> [(net_name, [non-R component pins])]

    with open(qcv_path) as f:
        for line in f:
            if not line.startswith("FlatNet:"):
                continue
            # FlatNet: 'NET_NAME' pin1 pin2 ...
            m = re.match(r"FlatNet:\s+'([^']+)'\s+(.*)", line.strip())
            if not m:
                continue
            net_name = m.group(1)
            pins = m.group(2).split()

            # Find any R### pins on this net
            r_pins_on_net = [p for p in pins if re.match(r'R\d+-\d+$', p)]
            non_r_pins = [p for p in pins if not re.match(r'R\d+-\d+$', p)]

            for rp in r_pins_on_net:
                r_refdes = rp.split("-")[0]  # "R426-2" -> "R426"
                r_nets.setdefault(r_refdes, []).append((net_name, non_r_pins))

    return r_nets


def resolve_via_pin(pin_str, r_net_map, signal_net_name):
    """Resolve 'LVC8T245 via R426' or 'HVD75 R via R119' to actual pin.

    Looks up the series resistor in the netlist map and returns the component
    pin on the far side (the net that is NOT the signal net).

    For 'U77-31 via R260' (already has specific pin), strips the 'via Rxxx'.
    """
    via_match = re.match(r'(.+?)\s+via\s+(R\d+)$', pin_str)
    if not via_match:
        return pin_str  # no "via" pattern, return as-is

    prefix = via_match.group(1)
    r_refdes = via_match.group(2)

    # If prefix is already a specific pin (e.g., "U77-31"), just return it
    if re.match(r'U\d+-\w+$', prefix):
        return prefix

    # Look up far-side pin from netlist
    if r_refdes not in r_net_map:
        return pin_str  # can't resolve, return original

    nets = r_net_map[r_refdes]
    # Find the net that is NOT the signal net (or any known signal net)
    for net_name, non_r_pins in nets:
        if net_name != signal_net_name and non_r_pins:
            # Limit to at most 3 pins to avoid massive headings on buses
            if len(non_r_pins) > 3:
                return ", ".join(non_r_pins[:3]) + " ..."
            return ", ".join(non_r_pins)

    return pin_str  # fallback


def resolve_pins_in_row(driver_pin, rx_pins, r_net_map, signal_net_name):
    """Resolve both driver and receiver pin strings."""
    driver_resolved = resolve_via_pin(driver_pin, r_net_map, signal_net_name)
    rx_resolved = resolve_via_pin(rx_pins, r_net_map, signal_net_name)
    return driver_resolved, rx_resolved


# ---------------------------------------------------------------------------
# Document building helpers
# ---------------------------------------------------------------------------

def load_csv(path):
    """Load CSV preserving insertion order; group by base signal."""
    runs = OrderedDict()  # key = (net_name, run_id)
    with open(path) as f:
        for row in csv.DictReader(f):
            key = (row["net_name"], row["run_id"])
            runs[key] = row
    # Group into signal -> list of (run_id, row)
    signals = OrderedDict()
    for (net, rid), row in runs.items():
        signals.setdefault(net, []).append((rid, row))
    return signals


def get_tr_tf_override(run_data):
    """Determine which rise/fall time spec override applies, if any.

    Returns:
      "lvc_input"  — SN55HVD75 drives → LVC8T245 receives (33ns, datasheet)
      "rs422"      — LVC8T245 drives → SN55HVD75 receives (50ns, TIA-422-B)
      None         — use default JEDEC spec
    """
    drv = run_data.get("driver_model", "")
    rx = run_data.get("rx_model", "")
    if drv in HVD75_DRIVER_MODELS:
        # SN55 is driving — receiver-side spec (LVC8T245 input) governs
        return "lvc_input"
    if rx in HVD75_MODELS:
        # LVC8T245 (or other) is driving into SN55 — RS-422 spec governs
        return "rs422"
    return None


def is_hvd75_signal(run_data):
    """Check if a run case involves an SN55HVD75 (RS-422) model."""
    return get_tr_tf_override(run_data) is not None


def pass_fail(comp, measured, threshold):
    if comp == ">=":
        return "PASS" if measured >= threshold else "FAIL"
    elif comp == "<=":
        return "PASS" if measured <= threshold else "FAIL"
    return "PASS"


def build_param_rows(run_data, specs, vdd, tr_tf_override=None):
    """Build 7 parameter rows for one run case. Returns list of
    (param_name, measured_str, spec_str, pf_str).

    tr_tf_override: "lvc_input" | "rs422" | None
    """
    voh = float(run_data["voh"])
    vol = float(run_data["vol"])
    tr = float(run_data["rise_time_ns"])
    tf = float(run_data["fall_time_ns"])
    nmh = float(run_data["nmh"])
    nml = float(run_data["nml"])
    peak = float(run_data["peak_v"])
    trough = float(run_data["trough_v"])

    os_peak_ok = peak <= specs["os_peak"]
    os_trough_ok = trough >= specs["os_trough"]
    os_pf = "PASS" if (os_peak_ok and os_trough_ok) else "FAIL"
    os_measured = f"Peak={peak:.3f}V, Trough={trough:.3f}V"

    # Select rise/fall time spec based on signal direction
    if tr_tf_override == "lvc_input":
        tr_spec = LVC_INPUT_TR_TF_SPEC
        tf_spec = LVC_INPUT_TR_TF_SPEC
    elif tr_tf_override == "rs422":
        tr_spec = RS422_TR_TF_SPEC
        tf_spec = RS422_TR_TF_SPEC
    else:
        tr_spec = specs["tr"]
        tf_spec = specs["tf"]

    rows = [
        ("VOH (V)",                  f"{voh:.3f}",
         specs["voh"][0], pass_fail(specs["voh"][2], voh, specs["voh"][1])),
        ("VOL (V)",                  f"{vol:.3f}",
         specs["vol"][0], pass_fail(specs["vol"][2], vol, specs["vol"][1])),
        ("Rise Time (ns)",           f"{tr:.2f}",
         tr_spec[0],  pass_fail(tr_spec[2], tr, tr_spec[1])),
        ("Fall Time (ns)",           f"{tf:.2f}",
         tf_spec[0],  pass_fail(tf_spec[2], tf, tf_spec[1])),
        ("Noise Margin High (V)",    f"{nmh:.3f}",
         specs["nmh"][0], pass_fail(specs["nmh"][2], nmh, specs["nmh"][1])),
        ("Noise Margin Low (V)",     f"{nml:.3f}",
         specs["nml"][0], pass_fail(specs["nml"][2], nml, specs["nml"][1])),
        ("Overshoot/Undershoot (V)", os_measured,
         specs["os_spec"], os_pf),
    ]
    return rows


def set_cell(cell, text, bold=False, size_pt=8, align=None):
    """Set cell text with formatting, clearing existing content."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if align:
        para.alignment = align


def add_run_table(doc, sig_label, run_data, specs, vdd, r_net_map,
                  signal_net_name):
    """Add a single run's 7-row parameter table to the document."""
    headers = ["Signal", "Parameter", "Measured Value",
               "Specification", "Pass/Fail", "Driver Pin", "Receiver Pin"]

    override = get_tr_tf_override(run_data)
    param_rows = build_param_rows(run_data, specs, vdd,
                                  tr_tf_override=override)

    # Resolve pin names
    driver_raw = run_data["driver_pin"]
    rx_raw = run_data["rx_pins"]
    driver, rx = resolve_pins_in_row(driver_raw, rx_raw, r_net_map,
                                     signal_net_name)

    # Build table rows
    all_rows = []
    for param, meas, spec, pf in param_rows:
        all_rows.append((sig_label, param, meas, spec, pf, driver, rx))

    # Create table
    table = doc.add_table(rows=1 + len(all_rows), cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, h in enumerate(headers):
        set_cell(table.rows[0].cells[ci], h, bold=True, size_pt=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, (sig, param, meas, spec, pf, drv, rcv) in enumerate(all_rows):
        row = table.rows[ri + 1]
        set_cell(row.cells[0], sig, size_pt=8)
        set_cell(row.cells[1], param, size_pt=8)
        set_cell(row.cells[2], meas, size_pt=8,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[3], spec, size_pt=8)
        # Color-code pass/fail
        set_cell(row.cells[4], pf, size_pt=8, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        if pf == "FAIL":
            row.cells[4].paragraphs[0].runs[0].font.color.rgb = \
                RGBColor(0xCC, 0x00, 0x00)
        else:
            row.cells[4].paragraphs[0].runs[0].font.color.rgb = \
                RGBColor(0x00, 0x80, 0x00)
        set_cell(row.cells[5], drv, size_pt=8)
        set_cell(row.cells[6], rcv, size_pt=8)

    # Set column widths
    widths = [Inches(1.8), Inches(1.5), Inches(1.4),
              Inches(1.6), Inches(0.6), Inches(1.2), Inches(1.4)]
    for row in table.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = w

    # Spacing after table
    doc.add_paragraph("")

    return override


def build_document(csv_path, vdd, section_num, section_title, output_path,
                   qcv_path="temp.qcv"):
    """Build a complete restructured Word document."""
    specs = JEDEC_SPECS[vdd]
    signals = load_csv(csv_path)
    r_net_map = build_resistor_pin_map(qcv_path)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    # Section heading
    doc.add_heading(
        f"{section_num}    {section_title} ({specs['standard']})", level=1)

    # Brief spec summary paragraph
    if vdd == 3.3:
        doc.add_paragraph(
            "Per JESD8C.01: VIH ≥ 2.0V, VIL ≤ 0.8V, VOH ≥ 2.4V, "
            "VOL ≤ 0.4V, "
            "NM ≥ 0.4V, tr/tf ≤ 10ns @ 50pF, "
            "Absolute max: VDD+0.3V = 3.6V (overshoot), "
            "GND−0.3V = −0.3V (undershoot)."
        )
    else:
        doc.add_paragraph(
            "Per JESD8-7A: VIH ≥ 0.65×VDD = 1.17V, "
            "VIL ≤ 0.35×VDD = 0.63V, "
            "VOH ≥ VDD−0.45V = 1.35V, VOL ≤ 0.45V, "
            "NM ≥ 0.18V, tr/tf ≤ 10ns @ 50pF, "
            "Absolute max: VDD+0.3V = 2.1V (overshoot), "
            "GND−0.3V = −0.3V (undershoot)."
        )

    has_rs422_footnote = False
    has_lvc_input_footnote = False

    # Per-signal subsections
    for idx, (sig_name, run_cases) in enumerate(signals.items(), 1):
        has_multiple_runs = (len(run_cases) > 1 and
                            any(rid in ("A", "B") for rid, _ in run_cases))

        # Signal heading: e.g. "3.7.1  TR0_EXT_SYNC_P3V3"
        doc.add_heading(
            f"{section_num}.{idx}    {sig_name}", level=2)

        if has_multiple_runs:
            # Split into sub-subsections for each run
            for sub_idx, (run_id, run_data) in enumerate(run_cases, 1):
                # Resolve pins for the heading
                driver_raw = run_data["driver_pin"]
                rx_raw = run_data["rx_pins"]
                driver_res, rx_res = resolve_pins_in_row(
                    driver_raw, rx_raw, r_net_map, sig_name)

                # Sub-subsection heading: "3.7.1.1  U128-1 → U128-4"
                doc.add_heading(
                    f"{section_num}.{idx}.{sub_idx}    "
                    f"{driver_res} \u2192 {rx_res}",
                    level=3)

                override = add_run_table(
                    doc, sig_name, run_data, specs, vdd,
                    r_net_map, sig_name)
                if override == "rs422":
                    has_rs422_footnote = True
                elif override == "lvc_input":
                    has_lvc_input_footnote = True
        else:
            # Single run — table directly under the signal heading
            _, run_data = run_cases[0]
            override = add_run_table(
                doc, sig_name, run_data, specs, vdd,
                r_net_map, sig_name)
            if override == "rs422":
                has_rs422_footnote = True
            elif override == "lvc_input":
                has_lvc_input_footnote = True

    # Add footnotes at end of document
    if has_rs422_footnote or has_lvc_input_footnote:
        doc.add_paragraph("")
    if has_rs422_footnote:
        footnote_para = doc.add_paragraph()
        run = footnote_para.add_run(RS422_FOOTNOTE)
        run.font.size = Pt(8)
        run.italic = True
    if has_lvc_input_footnote:
        footnote_para = doc.add_paragraph()
        run = footnote_para.add_run(LVC_INPUT_FOOTNOTE)
        run.font.size = Pt(8)
        run.italic = True

    doc.save(output_path)
    n_signals = len(signals)
    n_runs = sum(len(rc) for rc in signals.values())
    print(f"Built {output_path}: {n_signals} signals, {n_runs} run tables")
    if has_rs422_footnote:
        print(f"  (RS-422 tr/tf exception: ≤ 50ns for LVC8T245 → SN55HVD75)")
    if has_lvc_input_footnote:
        print(f"  (LVC8T245 input tr/tf exception: ≤ 33ns for SN55HVD75 → LVC8T245)")


if __name__ == "__main__":
    print("=== Building P3V3 document ===")
    build_document(
        "p3v3_results.csv",
        3.3,
        "3.7",
        "3.3V LVCMOS Signal Test",
        "HyperLynx_SI_Analysis.docx",
    )

    print("\n=== Building P1V8 document ===")
    build_document(
        "p1v8_results.csv",
        1.8,
        "3.8",
        "1.8V LVCMOS Signal Test",
        "HyperLynx_P1V8_SI_Analysis.docx",
    )

    print("\nDone.")
